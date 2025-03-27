import os
import hashlib
from docx import Document

def save_transcription_to_txt(response, output_txt_path):
    
    text = response.text  # acesso direto ao atributo
    with open(output_txt_path, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"📄 Transcrição salva em: {output_txt_path}")

def buscar_pdfs_em_diretorio(diretorio):
    return [
        os.path.join(diretorio, nome)
        for nome in os.listdir(diretorio)
        if nome.lower().endswith(".pdf")
    ]

def gerar_hash(caminho):
    with open(caminho, "rb") as f:
        conteudo = f.read()
    return hashlib.md5(conteudo).hexdigest()

def buscar_not_pdfs_em_diretorio(diretorio):
    return [
        os.path.join(diretorio, nome)
        for nome in os.listdir(diretorio)
        if (nome.lower().endswith(".txt") or nome.lower().endswith(".doc") or nome.lower().endswith(".docx"))
    ]

def buscar_videos_em_diretorio(diretorio):
    return [
        os.path.join(diretorio, nome)
        for nome in os.listdir(diretorio)
        if (nome.lower().endswith(".mp4"))
    ]

def carregar_txt(caminho):
    with open(caminho, 'r', encoding='utf-8') as f:
        return f.read()

def carregar_docx(caminho):
    doc = Document(caminho)
    return "\n".join([p.text for p in doc.paragraphs])

def carregar_doc (caminho):
    doc = Document(caminho)
    return "\n".join([p.text for p in doc.paragraphs])
